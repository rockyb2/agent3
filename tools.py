from smolagents import Tool
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import cm
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from openpyxl.worksheet.table import Table, TableStyleInfo
from openpyxl.chart import BarChart, Reference
from openpyxl.drawing.image import Image

# from email.mime.text import MIMEText
# from email.mime.multipart import MIMEMultipart
# from email.mime.base import MIMEBase
# from email import encoders
# import smtplib
import sib_api_v3_sdk
from sib_api_v3_sdk.rest import ApiException
import base64
import os
import logging
from typing import Optional

logger = logging.getLogger(__name__)


class BuildExcelPro(Tool):
    name = "BuildExcelPro"
    description = (
        "CrÃ©e un fichier Excel professionnel avec tableau, formules, graphiques et image."
    )

    inputs = {
        "name": {
            "type": "string",
            "description": "Nom du fichier (sans .xlsx)",
            
        },
        "headers": {
            "type": "array",
            "description": "Liste des titres de colonnes",
            
        },
        "rows": {
            "type": "array",
            "description": "Toutes les donnÃ©es ligne par ligne",
            
        },
        # "chart": {
        #     "type": "boolean",
        #     "description": "Ajouter un graphique ?",
            
        # },
        # "image_path": {
        #     "type": "string",
        #     "description": "Chemin d'une image Ã  insÃ©rer",
            
        # }
    }

    output_type = "string"

    def forward(self, name, headers, rows, ):
        try:
            wb = Workbook()
            ws = wb.active
            ws.title = "Feuille1"

            # ---- EN-TÃŠTES ----
            ws.append(headers)
            for cell in ws[1]:
                cell.font = Font(bold=True, color="FFFFFF")
                cell.fill = PatternFill("solid", fgColor="4472C4")
                cell.alignment = Alignment(horizontal="center")

            # ---- LIGNES ----
            for row in rows:
                ws.append(row)

            # ---- TABLE STYLE ----
            last_row = ws.max_row
            last_col = ws.max_column
            table = Table(
                displayName="Table1",
                ref=f"A1:{chr(64 + last_col)}{last_row}"
            )
            style = TableStyleInfo(
                name="TableStyleMedium9",
                showFirstColumn=False,
                showLastColumn=False,
                showRowStripes=True,
                showColumnStripes=False
            )
            table.tableStyleInfo = style
            ws.add_table(table)

            # ---- AUTO-WIDTH ----
            for col in ws.columns:
                max_length = max(len(str(c.value)) for c in col)
                ws.column_dimensions[col[0].column_letter].width = max_length + 2

            # ---- GRAPHIQUE ----
            # if chart:
            #     chart_obj = BarChart()
            #     chart_obj.title = "Graphique des donnÃ©es"

            #     data = Reference(ws, min_col=2, min_row=1,
            #                     max_row=last_row, max_col=last_col)
            #     chart_obj.add_data(data, titles_from_data=True)

            #     cats = Reference(ws, min_col=1, min_row=2, max_row=last_row)
            #     chart_obj.set_categories(cats)

            #     ws.add_chart(chart_obj, f"{chr(66 + last_col)}2")

            

            # ---- SAUVEGARDE ----
            file_name = f"{name}.xlsx"
            wb.save(file_name)

            return f"Excel '{file_name}' gÃ©nÃ©rÃ© avec succÃ¨s !||{file_name}"

        except Exception as e:
            return f"Erreur Excel Pro : {e}"


class BuildPDF(Tool):
    name = "BuildPDF"
    description = "GÃ©nÃ¨re un PDF professionnel avec titre, contenu, marges et styles optimisÃ©s."

    inputs = {
        "name": {"type": "string", "description": "Nom du fichier PDF sans extension"},
        "title": {"type": "string", "description": "Titre du PDF"},
        "content": {"type": "string", "description": "Texte du PDF"},
    }
    output_type = "string"

    def forward(self, name: str, title: str, content: str) -> str:
        try:
            file_name = f"{name}.pdf"
            styles = getSampleStyleSheet()
            style_title = styles["Title"]
            style_body = styles["BodyText"]

            doc = SimpleDocTemplate(
                file_name,
                pagesize=A4,
                leftMargin=2*cm,
                rightMargin=2*cm,
                topMargin=2*cm,
                bottomMargin=2*cm
            )

            story = [
                Paragraph(title, style_title),
                Spacer(1, 12),
                Paragraph(content.replace("\n", "<br/>"), style_body)
            ]

            doc.build(story)

            return f"PDF '{file_name}' gÃ©nÃ©rÃ© avec succÃ¨s||{file_name}"

        except Exception as e:
            return f"Erreur PDF : {str(e)}"


class BuildWord(Tool):
    name = "BuildWord"
    description = (
        "CrÃ©e un document Word professionnel (.docx) Ã  partir d'un contenu structurÃ©. "
        "Utilise cette outil pour gÃ©nÃ©rer des lettres de motivation, CV, rapports, etc. "
        "Le rÃ©sultat sera un fichier Word avec mise en page propre (marges, police, interligne)."
    )

    inputs = {
        "title": {"type": "string", "description": "Titre principal du document (ex: 'Lettre de Motivation')"},
        "recipient": {"type": "string", "description": "Nom et adresse du destinataire (multiligne possible)"},
        "sender": {"type": "string", "description": "Vos coordonnÃ©es (nom, adresse, email, tel)"},
        "date": {"type": "string", "description": "Date (ex: 'Paris, le 23 dÃ©cembre 2025')"},
        "subject": {"type": "string", "description": "Objet de la lettre"},
        "body": {"type": "string", "description": "Corps complet de la lettre en texte brut (paragraphes sÃ©parÃ©s par \\n\\n). Tu peux inclure des listes avec - item"},
        "filename": {"type": "string", "description": "Nom du fichier sans extension (ex: 'Lettre_Motivation_Jonathan')"}
    }

    output_type = "string"

    def forward(self, title: str, recipient: str, sender: str, date: str, subject: str, body: str, filename: str):
        try:
            doc = Document()

            # === Configuration page professionnelle (format lettre franÃ§aise) ===
            section = doc.sections[0]
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

            # === Style global ===
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Arial'
            font.size = Pt(11)

            # === En-tÃªte : Vos coordonnÃ©es (alignÃ© Ã  droite) ===
            p = doc.add_paragraph(sender)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # === Destinataire (alignÃ© Ã  gauche) ===
            doc.add_paragraph(recipient)

            # === Date (alignÃ© Ã  gauche ou droite selon norme) ===
            doc.add_paragraph(date)

            # === Objet ===
            p = doc.add_paragraph()
            p.add_run('Objet : ').bold = True
            p.add_run(subject)

            doc.add_paragraph()  # Espace

            # === Titre centrÃ© ===
            title_p = doc.add_paragraph(title)
            title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_p.runs[0]
            title_run.font.size = Pt(16)
            title_run.bold = True

            doc.add_paragraph()  # Espace

            # === Corps de la lettre ===
            paragraphs = body.split('\n\n')
            for para_text in paragraphs:
                para_text = para_text.strip()
                if para_text.startswith('- '):
                    # Liste Ã  puces
                    items = para_text.split('\n- ')
                    for i, item in enumerate(items):
                        if i == 0 and not item.startswith('- '):
                            item = item[2:] if item.startswith('- ') else item
                        p = doc.add_paragraph(item.strip('- '), style='List Bullet')
                else:
                    doc.add_paragraph(para_text)

            # === Formule de politesse ===
            doc.add_paragraph()
            doc.add_paragraph("Je vous prie dâ€™agrÃ©er, Madame, Monsieur, lâ€™expression de mes salutations distinguÃ©es.")

            doc.add_paragraph()  # Espace signature
            doc.add_paragraph("Jonathan Zadi")

            # === Sauvegarde ===
            safe_filename = "".join(c for c in filename if c.isalnum() or c in (' ', '-', '_')).rstrip()
            file_path = f"{safe_filename}.docx"
            doc.save(file_path)

            return f"Fichier Word professionnel crÃ©Ã© avec succÃ¨s : {os.path.abspath(file_path)} || {os.path.abspath(file_path)}"

        except Exception as e:
            return f"Erreur lors de la crÃ©ation du document Word : {str(e)}"


# class SendMail(Tool):
#     name = "send_mail"
#     description = (
#         "Envoie un e-mail avec support HTML et piÃ¨ces jointes optionnelles. "
#         "Inputs: smtp_server, smtp_port, sender_email, sender_password, "
#         "recipient_email, subject, message, is_html (optionnel), attachment_path (optionnel)"
#     )

#     inputs = {
#         "smtp_server": {"type": "string", "description": "Adresse du serveur SMTP (ex: smtp.gmail.com)"},
#         "smtp_port": {"type": "number", "description": "Port SMTP (ex: 587 pour TLS, 465 pour SSL)"},
#         "sender_email": {"type": "string", "description": "Adresse email expÃ©diteur (ex: davjonathan6@gmail.com)"},
#         "sender_password": {"type": "string", "description": "Mot de passe ou App Password de l'expÃ©diteur (ex: qbcqkupoknwgeenf)"},
#         "recipient_email": {"type": "string", "description": "Adresse email destinataire (peut Ãªtre une liste sÃ©parÃ©e par des virgules)"},
#         "subject": {"type": "string", "description": "Sujet du mail"},
#         "message": {"type": "string", "description": "Contenu du mail (texte ou HTML si is_html=True)"},
#         "is_html": {"type": "boolean", "description": "Si True, le message est interprÃ©tÃ© comme HTML (dÃ©faut: False)", "nullable": True},
#         "attachment_path": {"type": "string", "description": "Chemin vers un fichier Ã  joindre (optionnel)", "nullable": True}
#     }

#     output_type = "string"

#     def forward(
#         self,
#         smtp_server: str,
#         smtp_port: int,
#         sender_email: str,
#         sender_password: str,
#         recipient_email: str,
#         subject: str,
#         message: str,
#         is_html: bool = False,
#         attachment_path: Optional[str] = None
#     ) -> str:
#         try:
#             # Validation
#             if not all([smtp_server, sender_email, sender_password, recipient_email, subject, message]):
#                 return "Erreur: Tous les champs obligatoires doivent Ãªtre remplis."

#             # CrÃ©ation du mail
#             msg = MIMEMultipart()
#             msg["From"] = sender_email
#             msg["To"] = recipient_email
#             msg["Subject"] = subject

#             # Ajouter le message (texte ou HTML)
#             msg_type = "html" if is_html else "plain"
#             msg.attach(MIMEText(message, msg_type))

#             # Ajouter la piÃ¨ce jointe si fournie
#             if attachment_path and os.path.isfile(attachment_path):
#                 with open(attachment_path, "rb") as attachment:
#                     part = MIMEBase('application', 'octet-stream')
#                     part.set_payload(attachment.read())
#                     encoders.encode_base64(part)
#                     part.add_header(
#                         'Content-Disposition',
#                         f'attachment; filename= {os.path.basename(attachment_path)}'
#                     )
#                     msg.attach(part)
#                 logger.info(f"PiÃ¨ce jointe ajoutÃ©e: {attachment_path}")

#             # Connexion SMTP
#             if smtp_port == 465:
#                 # SSL
#                 server = smtplib.SMTP_SSL(smtp_server, smtp_port)
#             else:
#                 # TLS
#                 server = smtplib.SMTP(smtp_server, smtp_port)
#                 server.starttls()

#             server.login(sender_email, sender_password)

#             # Envoi
#             server.send_message(msg)
#             server.quit()

#             logger.info(f"Email envoyÃ© de {sender_email} Ã  {recipient_email}")
#             attachment_info = f" avec piÃ¨ce jointe ({os.path.basename(attachment_path)})" if attachment_path else ""
#             return f"E-mail envoyÃ© Ã  {recipient_email} avec succÃ¨s !{attachment_info}"

#         except FileNotFoundError:
#             return f"Erreur: Le fichier de piÃ¨ce jointe '{attachment_path}' n'existe pas."
#         except smtplib.SMTPAuthenticationError:
#             return "Erreur: Ã‰chec de l'authentification. VÃ©rifiez l'email et le mot de passe."
#         except smtplib.SMTPException as e:
#             return f"Erreur SMTP lors de l'envoi du mail : {e}"
#         except Exception as e:
#             logger.error(f"Erreur envoi email: {e}")
#             return f"Erreur lors de l'envoi du mail : {e}"

class SendMail(Tool):
    name = "send_mail"
    description = (
        "Envoie un email via Brevo (Sendinblue) avec support HTML "
        "et piÃ¨ce jointe optionnelle."
    )

    inputs = {
        "recipient_email": {"type": "string", "description": "Email du destinataire"},
        "subject": {"type": "string", "description": "Sujet de l'email"},
        "message": {"type": "string", "description": "Contenu du message (HTML ou texte)"},
        "is_html": {"type": "boolean", "description": "Message HTML ?", "nullable": True},
        "attachment_path": {"type": "string", "description": "Chemin du fichier joint", "nullable": True}
    }

    output_type = "string"

    def forward(
        self,
        recipient_email: str,
        subject: str,
        message: str,
        is_html: bool = False,
        attachment_path: Optional[str] = None
    ) -> str:

        try:
            api_key = os.getenv("BREVO_API_KEY")
            sender_email = os.getenv("SENDER_EMAIL", "no-reply@agent-ia.com")

            if not api_key:
                return "Erreur : clÃ© API Brevo manquante."

            configuration = sib_api_v3_sdk.Configuration()
            configuration.api_key['api-key'] = api_key

            api_instance = sib_api_v3_sdk.TransactionalEmailsApi(
                sib_api_v3_sdk.ApiClient(configuration)
            )

            email_data = {
                "to": [{"email": recipient_email}],
                "subject": subject,
                "sender": {"email": sender_email, "name": "candidAI"},
                "htmlContent": message if is_html else None,
                "textContent": message if not is_html else None,
            }

            # ðŸ“Ž PiÃ¨ce jointe
            if attachment_path and os.path.isfile(attachment_path):
                with open(attachment_path, "rb") as f:
                    encoded = base64.b64encode(f.read()).decode()

                email_data["attachment"] = [{
                    "content": encoded,
                    "name": os.path.basename(attachment_path)
                }]

            api_instance.send_transac_email(email_data)

            return f"ðŸ“§ Email envoyÃ© avec succÃ¨s Ã  {recipient_email}"

        except ApiException as e:
            return f"Erreur Brevo API : {e}"
        except Exception as e:
            return f"Erreur lors de l'envoi de l'email : {e}"
